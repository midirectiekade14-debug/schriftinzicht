"""
Re-parse all Smijtegelt DOCX files with correct sermon boundaries.
Deletes old data and reloads clean sermons.
"""
import re, json, os, sys
from docx import Document
from dotenv import load_dotenv
import requests

load_dotenv(os.path.join(os.path.expanduser("~"), "schriftinzicht", ".env"))

BASE = os.path.join(os.path.expanduser("~"), "schriftinzicht")
AUTHOR_ID = 7  # Smijtegelt

URL = os.getenv("SUPABASE_URL") + "/rest/v1"
KEY = os.getenv("SUPABASE_SERVICE_KEY")
if not KEY:
    print("ERROR: SUPABASE_SERVICE_KEY not set")
    sys.exit(1)
HEADERS = {
    "apikey": KEY,
    "Authorization": f"Bearer {KEY}",
    "Content-Type": "application/json",
    "Prefer": "return=minimal"
}

# Load lookups
with open(os.path.join(BASE, "verse_lookup.json")) as f:
    VERSE_LOOKUP = json.load(f)
with open(os.path.join(BASE, "bible_books.json")) as f:
    books = json.load(f)
BOOK_LOOKUP = {}
for b in books:
    BOOK_LOOKUP[b["name"]] = b["id"]
    BOOK_LOOKUP[b["abbreviation"]] = b["id"]
# Common abbreviations
ABBREV = {
    "Gen": "Genesis", "Ex": "Exodus", "Lev": "Leviticus",
    "Num": "Numeri", "Deut": "Deuteronomium", "Joz": "Jozua",
    "Richt": "Richteren", "1 Sam": "1 Samuel", "2 Sam": "2 Samuel",
    "1 Kon": "1 Koningen", "2 Kon": "2 Koningen",
    "1 Kron": "1 Kronieken", "2 Kron": "2 Kronieken",
    "Neh": "Nehemia", "Est": "Esther", "Ps": "Psalmen",
    "Spr": "Spreuken", "Pred": "Prediker", "Hoogl": "Hooglied",
    "Jes": "Jesaja", "Jer": "Jeremia", "Klaagl": "Klaagliederen",
    "Ez": "Ezechiël", "Dan": "Daniël", "Hos": "Hosea",
    "Jo\u00ebl": "Joël", "Am": "Amos", "Ob": "Obadja", "Mi": "Micha",
    "Nah": "Nahum", "Hab": "Habakuk", "Zef": "Zefanja",
    "Hag": "Haggaï", "Zach": "Zacharia", "Mal": "Maleachi",
    "Matt": "Mattheüs", "Matth": "Mattheüs", "Mark": "Markus",
    "Luk": "Lukas", "Joh": "Johannes", "Hand": "Handelingen",
    "Rom": "Romeinen", "1 Kor": "1 Korinthe", "2 Kor": "2 Korinthe",
    "Gal": "Galaten", "Ef": "Efeze", "Fil": "Filippenzen",
    "Kol": "Kolossenzen", "1 Thess": "1 Thessalonicenzen",
    "2 Thess": "2 Thessalonicenzen", "1 Tim": "1 Timotheüs",
    "2 Tim": "2 Timotheüs", "Tit": "Titus", "Filem": "Filemon",
    "Hebr": "Hebreeën", "Jak": "Jakobus",
    "1 Petr": "1 Petrus", "2 Petr": "2 Petrus",
    "1 Joh": "1 Johannes", "2 Joh": "2 Johannes",
    "3 Joh": "3 Johannes", "Jud": "Judas", "Openb": "Openbaring van Johannes",
}
for abbr, full in ABBREV.items():
    if full in BOOK_LOOKUP:
        BOOK_LOOKUP[abbr] = BOOK_LOOKUP[full]

# Verse reference regex
VERSE_RE = re.compile(
    r'(Genesis|Exodus|Leviticus|Numeri|Deuteronomium|Jozua|Richteren|Ruth|'
    r'(?:1|2)\s*Samu[eë]l|(?:1|2)\s*Koningen|(?:1|2)\s*Kronieken|'
    r'Ezra|Nehemia|Esther|Job|Psalm(?:en)?|Spreuken|Prediker|Hooglied|'
    r'Jesaja|Jeremia|Klaagliederen|Ezech?i[eë]l|Dani[eë]l|'
    r'Hosea|Jo[eë]l|Amos|Obadja|Jona|Micha|Nahum|Habakuk|'
    r'Zefanja|Hagga[ïi]|Zacharia|Maleachi|'
    r'Matth?[eë\u00eb][uü\u00fcs]+|Markus|Lukas|Johannes|Handelingen|'
    r'Romeinen|(?:1|2)\s*Korinth[eë]|Galaten|Ef[eé]ze|Filippenzen|Kolossenzen|'
    r'(?:1|2)\s*Thessalonicenzen|(?:1|2)\s*Timoth[eë][uü]s|'
    r'Titus|Filemon|Hebree[eë]n|Jakobus|(?:1|2)\s*Petrus|'
    r'(?:1|2|3)\s*Johannes|Judas|Openbaring)'
    r'\s*(\d+)\s*[,:]\s*(\d+)',
    re.IGNORECASE
)


def extract_verse_ref(text):
    """Extract first verse reference. Returns (book, chapter, verse) or None."""
    m = VERSE_RE.search(text)
    if not m:
        return None
    book = m.group(1).strip()
    chapter = int(m.group(2))
    verse = int(m.group(3))
    return book, chapter, verse


def resolve_verse_id(book_name, chapter, verse):
    """Resolve to verse_id via lookups."""
    book_id = BOOK_LOOKUP.get(book_name)
    if not book_id:
        # Try case variations
        for k, v in BOOK_LOOKUP.items():
            if k.lower() == book_name.lower():
                book_id = v
                break
    if not book_id:
        return None
    key = f"{book_id}_{chapter}_{verse}"
    return VERSE_LOOKUP.get(key)


def split_on_pattern(paragraphs, pattern):
    """Split paragraphs into sermons using a regex pattern for headers."""
    sermons = []
    current_title = None
    current_lines = []

    for txt in paragraphs:
        if pattern.match(txt):
            if current_title and current_lines:
                full_text = '\n'.join(current_lines).strip()
                if len(full_text.split()) > 200:  # Skip TOC entries and tiny fragments
                    sermons.append({'title': current_title, 'text': full_text})
            current_title = txt.rstrip('.').strip()
            current_lines = []
        elif current_title is not None:
            if txt:
                current_lines.append(txt)

    if current_title and current_lines:
        full_text = '\n'.join(current_lines).strip()
        if len(full_text) > 100:
            sermons.append({'title': current_title, 'text': full_text})

    return sermons


def parse_16_predicaties():
    doc = Document(os.path.join(BASE, 'smijtegelt', '16_predicaties.docx'))
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    pattern = re.compile(
        r'^(Eerste|Tweede|Derde|Vierde|Vijfde|Zesde|Zevende|Achtste|Negende|'
        r'Tiende|Elfde|Twaalfde|Dertiende|Veertiende|Viertiende|Vijftiende|Zestiende)\s+'
        r'predicatie\s+over\s+', re.IGNORECASE)
    return split_on_pattern(paragraphs, pattern)


def parse_50_keurstoffen():
    doc = Document(os.path.join(BASE, 'smijtegelt', '50_keurstoffen.docx'))
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    pattern = re.compile(r'^[A-Z]+E?\s+PREDICATIE[,.]?\s+[Oo][Vv][Ee][Rr]\s+')
    return split_on_pattern(paragraphs, pattern)


def parse_52_catechismus():
    doc = Document(os.path.join(BASE, 'smijtegelt', '52_catechismus.docx'))
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    pattern = re.compile(r'^ZONDAG\s+(\d+)', re.IGNORECASE)
    sermons = []
    current_title = None
    current_lines = []
    for txt in paragraphs:
        m = pattern.match(txt)
        if m:
            if current_title and current_lines:
                full_text = '\n'.join(current_lines).strip()
                if len(full_text) > 100:
                    sermons.append({'title': current_title, 'text': full_text})
            current_title = f"Zondag {m.group(1)}"
            current_lines = []
        elif current_title is not None:
            if txt:
                current_lines.append(txt)
    if current_title and current_lines:
        full_text = '\n'.join(current_lines).strip()
        if len(full_text) > 100:
            sermons.append({'title': current_title, 'text': full_text})
    return sermons


def parse_deel(filename, collection_suffix=""):
    """Parse Een Woord op Zijn Tijd deel files.
    Headers contain 'OVER' + verse reference and are actual content markers."""
    doc = Document(os.path.join(BASE, 'smijtegelt', filename))

    # Find sermon headers: "PREEK OVER <verse>" (not just "EERSTE PREEK" from TOC)
    pattern = re.compile(r'^.+PREEK\s+OVER\s+', re.IGNORECASE)
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    return split_on_pattern(paragraphs, pattern)


def parse_zestal():
    doc = Document(os.path.join(BASE, 'smijtegelt', 'zestal.docx'))
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    pattern = re.compile(
        r'^(Eerste|Tweede|Derde|Vierde|Vijfde|Zesde)\s+predicatie\s+over\s+',
        re.IGNORECASE)
    return split_on_pattern(paragraphs, pattern)


def load_collection(collection, sermons):
    """Delete old data and load new sermons for a collection."""
    print(f"\n{'='*60}")
    print(f"{collection}: {len(sermons)} preken gevonden")
    print(f"{'='*60}")

    if not sermons:
        print("  Geen preken, overslaan.")
        return

    # Show first few titles
    for s in sermons[:5]:
        wc = len(s['text'].split())
        print(f"  {wc:>6} w | {s['title'][:70]}")
    if len(sermons) > 5:
        print(f"  ... en nog {len(sermons)-5} meer")

    # Delete old
    r = requests.delete(
        f"{URL}/sermons?author_id=eq.{AUTHOR_ID}&source_collection=eq.{collection}",
        headers=HEADERS
    )
    print(f"  Oude data verwijderd: {r.status_code}")

    # Build batch
    batch = []
    for s in sermons:
        vref = extract_verse_ref(s['title']) or extract_verse_ref(s['text'][:500])
        verse_id = None
        if vref:
            verse_id = resolve_verse_id(vref[0], vref[1], vref[2])

        batch.append({
            "author_id": AUTHOR_ID,
            "title": s['title'],
            "start_verse_id": verse_id,
            "sermon_text": s['text'],
            "source_collection": collection,
            "language": "nl",
            "word_count": len(s['text'].split()),
        })

    # Insert
    inserted = 0
    for i in range(0, len(batch), 500):
        chunk = batch[i:i+500]
        r = requests.post(f"{URL}/sermons", headers=HEADERS, json=chunk)
        if r.status_code in (200, 201):
            inserted += len(chunk)
        else:
            print(f"  Fout batch {i}: {r.status_code} {r.text[:200]}")

    print(f"  Geladen: {inserted}/{len(batch)}")


def main():
    # Parse all collections
    collections = [
        ('16 Predicaties', parse_16_predicaties),
        ('50 Keurstoffen', parse_50_keurstoffen),
        ('52 Catechismuspreken', parse_52_catechismus),
        ('Een Woord op Zijn Tijd deel 1', lambda: parse_deel('deel1.docx')),
        ('Een Woord op Zijn Tijd deel 2', lambda: parse_deel('deel2.docx')),
        ('Een Woord op Zijn Tijd deel 3', lambda: parse_deel('deel3.docx')),
        ('Een Woord op Zijn Tijd deel 4', lambda: parse_deel('deel4.docx')),
        ('Zestal Leerredenen', parse_zestal),
    ]

    for collection, parser_fn in collections:
        sermons = parser_fn()
        load_collection(collection, sermons)

    # Clean up old "Een Woord op Zijn Tijd" (without deel) entries
    for old_coll in ['Een Woord op Zijn Tijd', 'Een Woord op Zijn Tijd deel 2',
                     'Een Woord op Zijn Tijd deel 3', 'Een Woord op Zijn Tijd deel 4']:
        r = requests.delete(
            f"{URL}/sermons?author_id=eq.{AUTHOR_ID}&source_collection=eq.{old_coll}",
            headers=HEADERS
        )
        if r.status_code == 204:
            print(f"\nOude collectie '{old_coll}' verwijderd")

    print("\nKlaar!")


if __name__ == "__main__":
    main()
