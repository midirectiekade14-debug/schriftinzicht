#!/usr/bin/env python3
"""
Parse Calvijn bijbelverklaring files from theologienet.nl into JSON.
- Genesis 1-3: from DOCX (text-based)
- Romeinen - 2 Korinthe, Galaten - Filemon, Hebreeen - Judas: from PDF (scanned, OCR via Tesseract)

Output format per entry:
{
    "book": "Genesis",
    "chapter": 1,
    "verse": 1,
    "verse_end": null,
    "text": "commentary text..."
}
"""

import json
import re
import sys
import io
import os
import time
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

BASE_DIR = Path("C:/Users/midir/schriftinzicht/scraped/calvijn")
OUTPUT_FILE = Path("C:/Users/midir/schriftinzicht/scraped/calvijn_nl_theologienet.json")

# Book name normalization
BOOK_NAMES = {
    "genesis": "Genesis",
    "gen": "Genesis",
    "romeinen": "Romeinen",
    "rom": "Romeinen",
    "1 korinthe": "1 Korinthe",
    "1 kor": "1 Korinthe",
    "1 korinthiërs": "1 Korinthe",
    "1 corinthe": "1 Korinthe",
    "2 korinthe": "2 Korinthe",
    "2 kor": "2 Korinthe",
    "2 korinthiërs": "2 Korinthe",
    "2 corinthe": "2 Korinthe",
    "galaten": "Galaten",
    "gal": "Galaten",
    "efeze": "Efeze",
    "efeziers": "Efeze",
    "efeziërs": "Efeze",
    "filippenzen": "Filippenzen",
    "fil": "Filippenzen",
    "kolossenzen": "Kolossenzen",
    "kol": "Kolossenzen",
    "1 thessalonicenzen": "1 Thessalonicenzen",
    "1 thess": "1 Thessalonicenzen",
    "2 thessalonicenzen": "2 Thessalonicenzen",
    "2 thess": "2 Thessalonicenzen",
    "1 timotheüs": "1 Timotheus",
    "1 timotheus": "1 Timotheus",
    "1 tim": "1 Timotheus",
    "2 timotheüs": "2 Timotheus",
    "2 timotheus": "2 Timotheus",
    "2 tim": "2 Timotheus",
    "titus": "Titus",
    "tit": "Titus",
    "filemon": "Filemon",
    "filem": "Filemon",
    "hebreeën": "Hebreeen",
    "hebreeen": "Hebreeen",
    "hebreen": "Hebreeen",
    "hebr": "Hebreeen",
    "jakobus": "Jakobus",
    "jak": "Jakobus",
    "1 petrus": "1 Petrus",
    "1 petr": "1 Petrus",
    "2 petrus": "2 Petrus",
    "2 petr": "2 Petrus",
    "1 johannes": "1 Johannes",
    "1 joh": "1 Johannes",
    "judas": "Judas",
    "jud": "Judas",
}


def parse_genesis_docx():
    """Parse Genesis 1-3 from DOCX file."""
    from docx import Document

    doc = Document(str(BASE_DIR / "calvijn-genesis-1-3.docx"))
    entries = []
    current_entry = None

    # Find where commentary begins (after the verse text blocks)
    # Commentary sections start with verse references like "1. In den beginne."
    # followed by explanatory text
    commentary_started = False
    in_verse_block = False

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        # Detect chapter headings
        chapter_match = re.match(r'^(\d+)(?:ste|de|e)?\s+HOOFDSTUK', text, re.IGNORECASE)
        if chapter_match:
            continue

        # Detect verse reference headers in commentary (e.g., "1. In den beginne." or "3. En God zeide.")
        # These appear after the block of all verses
        verse_header = re.match(r'^(\d+)\.\s+(.+)', text)

        # Check if we're in the commentary section (after paragraph index ~220 for Genesis)
        if i >= 220 and verse_header:
            # Save previous entry
            if current_entry and current_entry["text"].strip():
                entries.append(current_entry)

            verse_num = int(verse_header.group(1))
            # Determine chapter based on verse numbering and position
            chapter = _get_genesis_chapter(i, verse_num, doc.paragraphs)

            # Check for verse range (e.g., "6-8" or "9, 10")
            verse_end = None
            header_text = verse_header.group(0)
            range_match = re.match(r'^(\d+)[-–](\d+)', header_text)
            if range_match:
                verse_end = int(range_match.group(2))

            current_entry = {
                "book": "Genesis",
                "chapter": chapter,
                "verse": verse_num,
                "verse_end": verse_end,
                "text": ""
            }
            commentary_started = True
            continue

        # Accumulate commentary text
        if commentary_started and current_entry:
            # Skip if it looks like a new chapter heading
            if re.match(r'^\d+(?:ste|de|e)?\s+HOOFDSTUK', text, re.IGNORECASE):
                continue
            # Skip pure verse text blocks
            if re.match(r'^\d+\.\s+En\s+(God|Hij|de|het|zij|dat)\s', text) and len(text) < 300:
                # This might be a verse quote at start of commentary - include it
                pass
            current_entry["text"] += text + "\n"

    # Don't forget last entry
    if current_entry and current_entry["text"].strip():
        entries.append(current_entry)

    return entries


def _get_genesis_chapter(para_idx, verse_num, paragraphs):
    """Determine Genesis chapter based on context."""
    # Look backwards for chapter heading
    for j in range(para_idx - 1, max(0, para_idx - 50), -1):
        t = paragraphs[j].text.strip()
        ch_match = re.match(r'^(\d+)(?:ste|de|e)?\s+HOOFDSTUK', t, re.IGNORECASE)
        if ch_match:
            return int(ch_match.group(1))
        # Also check for "HOOFDSTUK X" pattern
        ch_match2 = re.match(r'^HOOFDSTUK\s+(\d+)', t, re.IGNORECASE)
        if ch_match2:
            return int(ch_match2.group(1))
    return 1  # Default to chapter 1


def ocr_pdf(pdf_path):
    """OCR a scanned PDF and return full text."""
    import fitz
    import pytesseract
    from PIL import Image

    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

    doc = fitz.open(str(pdf_path))
    full_text = ""
    total = len(doc)

    print(f"  OCR processing {total} pages...")
    for i in range(total):
        page = doc[i]
        mat = fitz.Matrix(2, 2)  # 2x zoom
        pix = page.get_pixmap(matrix=mat)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        text = pytesseract.image_to_string(img, lang='nld')
        full_text += text + "\n"

        if (i + 1) % 20 == 0 or i == total - 1:
            print(f"    Page {i+1}/{total}")

    return full_text


def detect_book_from_header(text):
    """Try to detect book name from header text in OCR output."""
    text_lower = text.lower().strip()
    for key, value in BOOK_NAMES.items():
        if key in text_lower:
            return value
    return None


def parse_ocr_text(full_text, expected_books):
    """
    Parse OCR text into verse-level commentary entries.

    The format typically has headers like:
    HEBREEN 1:1  or  ROMEINEN 3:5  or  GALATEN 2:11
    followed by verse text and then commentary.

    Commentary sections often start with a verse number and quote,
    then explanatory text follows.
    """
    entries = []
    lines = full_text.split('\n')

    current_book = None
    current_chapter = None
    current_verse = None
    current_verse_end = None
    current_text = []

    # Pattern for page headers like "HEBREEN 1:1, 2" or "ROMEINEN 3:5"
    header_pattern = re.compile(
        r'^([A-Z]\s*[A-Za-z]+(?:\s+[A-Za-z]+)?)\s+(\d+)\s*:\s*(\d+)',
        re.IGNORECASE
    )

    # Pattern for verse commentary starts like "1. In den beginne" or "Vs. 1."
    verse_pattern = re.compile(
        r'^\s*(?:Vs\.?\s*)?(\d+)\.\s+(.{5,})',
    )

    # Pattern for verse range like "5-8." or "5, 6."
    verse_range_pattern = re.compile(
        r'^\s*(?:Vs\.?\s*)?(\d+)\s*[-–,]\s*(\d+)\.\s+(.{5,})',
    )

    # Chapter heading pattern
    chapter_pattern = re.compile(
        r'(?:HOOFDSTUK|HOOFD[-\s]?STUK|Hoofdstuk)\s+(\d+)',
        re.IGNORECASE
    )

    for line_idx, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue

        # Try to detect book from page headers
        header_match = header_pattern.match(stripped)
        if header_match:
            book_candidate = header_match.group(1).strip()
            book_name = detect_book_from_header(book_candidate)
            if book_name and book_name in expected_books:
                current_book = book_name
                ch = int(header_match.group(2))
                if ch != current_chapter:
                    current_chapter = ch
            continue  # Skip header lines

        # Chapter heading
        ch_match = chapter_pattern.search(stripped)
        if ch_match and len(stripped) < 80:
            new_ch = int(ch_match.group(1))
            if 1 <= new_ch <= 150:
                current_chapter = new_ch
            continue

        if not current_book:
            # Try to detect book from text
            for book in expected_books:
                if book.upper() in stripped.upper():
                    current_book = book
                    break
            continue

        if not current_chapter:
            current_chapter = 1

        # Check for verse range first
        vr_match = verse_range_pattern.match(stripped)
        if vr_match:
            # Save previous
            if current_verse is not None and current_text:
                entries.append({
                    "book": current_book,
                    "chapter": current_chapter,
                    "verse": current_verse,
                    "verse_end": current_verse_end,
                    "text": '\n'.join(current_text).strip()
                })
            current_verse = int(vr_match.group(1))
            current_verse_end = int(vr_match.group(2))
            current_text = [vr_match.group(3).strip()]
            continue

        # Check for new verse commentary
        v_match = verse_pattern.match(stripped)
        if v_match:
            new_verse = int(v_match.group(1))
            # Sanity check: verse should be reasonable
            if 1 <= new_verse <= 176:
                # Save previous
                if current_verse is not None and current_text:
                    entries.append({
                        "book": current_book,
                        "chapter": current_chapter,
                        "verse": current_verse,
                        "verse_end": current_verse_end,
                        "text": '\n'.join(current_text).strip()
                    })
                current_verse = new_verse
                current_verse_end = None
                current_text = [v_match.group(2).strip()]
                continue

        # Accumulate text for current verse
        if current_verse is not None:
            current_text.append(stripped)

    # Last entry
    if current_verse is not None and current_text:
        entries.append({
            "book": current_book,
            "chapter": current_chapter,
            "verse": current_verse,
            "verse_end": current_verse_end,
            "text": '\n'.join(current_text).strip()
        })

    return entries


def clean_text(text):
    """Clean up OCR artifacts and normalize text."""
    # Fix common OCR issues
    text = re.sub(r'\s*-\s*\n\s*', '', text)  # Rejoin hyphenated words
    text = re.sub(r'\n{3,}', '\n\n', text)  # Normalize newlines
    text = text.strip()
    return text


def main():
    all_entries = []

    # 1. Parse Genesis DOCX
    print("=" * 60)
    print("Parsing Genesis 1-3 (DOCX)...")
    genesis_entries = parse_genesis_docx()
    for e in genesis_entries:
        e["text"] = clean_text(e["text"])
    all_entries.extend(genesis_entries)
    print(f"  -> {len(genesis_entries)} entries from Genesis")

    # 2. OCR and parse PDFs
    pdf_configs = [
        {
            "file": "calvijn-bijbelverklaring-romeinen-1-2-korinthe.pdf",
            "books": ["Romeinen", "1 Korinthe", "2 Korinthe"],
        },
        {
            "file": "calvijn-bijbelverklaring-galaten-filemon.pdf",
            "books": ["Galaten", "Efeze", "Filippenzen", "Kolossenzen",
                       "1 Thessalonicenzen", "2 Thessalonicenzen",
                       "1 Timotheus", "2 Timotheus", "Titus", "Filemon"],
        },
        {
            "file": "calvijn-bijbelverklaring-hebreeen-judas.pdf",
            "books": ["Hebreeen", "Jakobus", "1 Petrus", "2 Petrus",
                       "1 Johannes", "Judas"],
        },
    ]

    for config in pdf_configs:
        pdf_path = BASE_DIR / config["file"]
        print("=" * 60)
        print(f"Processing {config['file']}...")
        print(f"  Expected books: {', '.join(config['books'])}")

        # Save OCR text to cache file for debugging
        cache_file = pdf_path.with_suffix('.txt')
        if cache_file.exists():
            print(f"  Using cached OCR text from {cache_file.name}")
            with open(cache_file, 'r', encoding='utf-8') as f:
                full_text = f.read()
        else:
            full_text = ocr_pdf(pdf_path)
            with open(cache_file, 'w', encoding='utf-8') as f:
                f.write(full_text)
            print(f"  OCR text cached to {cache_file.name}")

        entries = parse_ocr_text(full_text, config["books"])
        for e in entries:
            e["text"] = clean_text(e["text"])

        # Filter out entries with very short text (likely noise)
        entries = [e for e in entries if len(e["text"]) > 30]
        all_entries.extend(entries)

        # Report per book
        book_counts = {}
        for e in entries:
            book_counts[e["book"]] = book_counts.get(e["book"], 0) + 1
        for book, count in sorted(book_counts.items()):
            print(f"  -> {book}: {count} entries")
        print(f"  -> Total: {len(entries)} entries")

    # Write output
    print("=" * 60)
    print(f"Writing {len(all_entries)} total entries to {OUTPUT_FILE}")
    with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
        json.dump(all_entries, f, ensure_ascii=False, indent=2)

    # Summary
    print("\n=== SUMMARY ===")
    book_summary = {}
    for e in all_entries:
        book_summary[e["book"]] = book_summary.get(e["book"], 0) + 1
    for book, count in sorted(book_summary.items()):
        print(f"  {book}: {count} entries")
    print(f"  TOTAL: {len(all_entries)} entries")


if __name__ == "__main__":
    main()
