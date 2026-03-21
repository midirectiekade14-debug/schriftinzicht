"""
Parse theologienet DOCX files into structured JSON.
Handles: Van der Groe Toetssteen (3 delen), Smijtegelt preken (8 werken).
"""

import json
import os
import re
from docx import Document

BASE = os.path.expanduser("~/schriftinzicht/scraped")


def extract_text(doc):
    """Extract all paragraph text from a Document."""
    return [(p.text.strip(), p.style.name if p.style else "Normal") for p in doc.paragraphs if p.text.strip()]


def save_json(data, filename):
    path = os.path.join(BASE, filename)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"  Saved {filename}: {len(data)} items")
    return len(data)


# --- Van der Groe: Toetssteen ---
def parse_toetssteen(deel_nr):
    """Parse Toetssteen by splitting on Heading sections.

    Strategy: split on Heading 3+ sections (topic headings like 'VOORREDE',
    'ZONDAG 20', 'Van de vergeving der zonden', etc.) which provide
    meaningful theological sections. Skip meta-headings like title page.
    """
    filename = f"groe-toetssteen-deel-{deel_nr}.docx"
    doc = Document(os.path.join(BASE, filename))

    source = f"Van der Groe - Toetssteen deel {deel_nr}"
    sections = []
    current_title = None
    current_text = []

    # Skip patterns for meta/title headings we don't want as section boundaries
    skip_titles = {'TOETSSTEEN', 'INHOUD', 'BLADWIJZER'}

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        style = p.style.name if p.style else "Normal"

        # Split on Heading 1-6 styles (these mark real sections)
        is_section = (
            style.startswith('Heading') and
            len(text) < 200 and  # Not a long paragraph accidentally styled as heading
            text.upper() not in skip_titles and
            not text.startswith('Spelling')  # Skip "Spelling aangepast" meta
        )

        if is_section:
            if current_title and current_text:
                sections.append({
                    "title": current_title,
                    "text": "\n\n".join(current_text),
                    "source_collection": source
                })
            current_title = text
            current_text = []
        elif current_title is not None:
            current_text.append(text)

    # Last section
    if current_title and current_text:
        sections.append({
            "title": current_title,
            "text": "\n\n".join(current_text),
            "source_collection": source
        })

    # If no sections found, save as single item
    if not sections:
        all_text = "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        sections = [{
            "title": f"Toetssteen deel {deel_nr}",
            "text": all_text,
            "source_collection": source
        }]

    return save_json(sections, f"groe-toetssteen-deel-{deel_nr}.json")


# --- Smijtegelt: sermon splitter ---
def split_sermons(docx_file, source_name, sermon_patterns=None):
    """Generic sermon splitter for Smijtegelt works."""
    doc = Document(os.path.join(BASE, docx_file))
    paragraphs = extract_text(doc)

    if sermon_patterns is None:
        sermon_patterns = [
            # "Eerste predicatie over ..."
            re.compile(r'^(Eerste|Tweede|Derde|Vierde|Vijfde|Zesde|Zevende|Achtste|Negende|Tiende|'
                       r'Elfde|Twaalfde|Dertiende|Veertiende|Vijftiende|Zestiende|Zeventiende|'
                       r'Achttiende|Negentiende|Twintigste|Eenentwintigste|Twee.ntwintigste|'
                       r'Drie.ntwintigste|Vierentwintigste|Vijfentwintigste|Zesentwintigste|'
                       r'Zevenentwintigste|Achtentwintigste|Negenentwintigste|Dertigste|'
                       r'Eenendertigste|Twee.ndertigste|Drie.ndertigste|Vierendertigste|'
                       r'Vijfendertigste|Zesendertigste|Zevenendertigste|Achtendertigste|'
                       r'Negenendertigste|Veertigste|Eenenveertigste|Twee.nveertigste|'
                       r'Drie.nveertigste|Vierenveertigste|Vijfenveertigste|Zesenveertigste|'
                       r'Zevenenveertigste|Achtenveertigste|Negenenveertigste|Vijftigste|'
                       r'Eenenvijftigste|Twee.nvijftigste)\s+(predicatie|preek|leerrede)',
                       re.IGNORECASE),
            # "XXSTE PREEK over ..."
            re.compile(r'^(EERSTE|TWEEDE|DERDE|VIERDE|VIJFDE|ZESDE|ZEVENDE|ACHTSTE|NEGENDE|TIENDE|'
                       r'ELFDE|TWAALFDE|DERTIENDE|VEERTIENDE|VIERTIENDE|VIJFTIENDE|ZESTIENDE|ZEVENTIENDE|'
                       r'ACHTTIENDE|NEGENTIENDE|TWINTIGSTE|EENENTWINTIGSTE|TWEE.NTWINTIGSTE|'
                       r'DRIE.NTWINTIGSTE|VIERENTWINTIGSTE|VIJFENTWINTIGSTE|ZESENTWINTIGSTE|'
                       r'ZEVENENTWINTIGSTE|ACHTENTWINTIGSTE|NEGENENTWINTIGSTE|DERTIGSTE|'
                       r'EENENDERTIGSTE|TWEE.NDERTIGSTE|DRIE.NDERTIGSTE|VIERENDERTIGSTE|'
                       r'VIJFENDERTIGSTE|ZESENDERTIGSTE|ZEVENENDERTIGSTE|ACHTENDERTIGSTE|'
                       r'NEGENENDERTIGSTE|VEERTIGSTE|EENENVEERTIGSTE|TWEE.NVEERTIGSTE|'
                       r'DRIE.NVEERTIGSTE|VIERENVEERTIGSTE|VIJFENVEERTIGSTE|ZESENVEERTIGSTE|'
                       r'ZEVENENVEERTIGSTE|ACHTENVEERTIGSTE|NEGENENVEERTIGSTE|VIJFTIGSTE|'
                       r'EENENVIJFTIGSTE|TWEE.NVIJFTIGSTE)\s+(PREDICATIE|PREEK|PREDIKATIE|LEERREDE)',
                       re.IGNORECASE),
            # Numbered: "1e predicatie", "22e preek" etc
            re.compile(r'^\d+[eE]\s+(predicatie|preek|predikatie|leerrede)', re.IGNORECASE),
        ]

    sermons = []
    current_title = None
    current_text = []
    in_preamble = True  # Skip front matter

    for text, style in paragraphs:
        is_sermon_start = any(pat.search(text) for pat in sermon_patterns)

        if is_sermon_start:
            in_preamble = False
            if current_title and current_text:
                sermons.append({
                    "title": current_title,
                    "text": "\n\n".join(current_text),
                    "source_collection": source_name
                })
            current_title = text
            current_text = []
        elif not in_preamble and current_title is not None:
            current_text.append(text)

    # Last sermon
    if current_title and current_text:
        sermons.append({
            "title": current_title,
            "text": "\n\n".join(current_text),
            "source_collection": source_name
        })

    # Filter out TOC entries: if we have duplicates, keep only items with
    # substantial text (>500 chars). This handles DOCX files where the
    # table of contents repeats the same headings as short summaries.
    if len(sermons) > 0:
        avg_len = sum(len(s["text"]) for s in sermons) / len(sermons)
        if avg_len < 2000:
            # Probably all short — keep everything
            pass
        else:
            # Filter out suspiciously short items (TOC entries)
            min_len = 1000
            filtered = [s for s in sermons if len(s["text"]) >= min_len]
            if filtered:
                sermons = filtered

    return sermons


def parse_smytegelt_16():
    sermons = split_sermons(
        "smytegelt-16-predicaties.docx",
        "Smijtegelt - 16 Predicaties"
    )
    return save_json(sermons, "smytegelt-16-predicaties.json")


def parse_smytegelt_50():
    sermons = split_sermons(
        "smytegelt-50-keurstoffen-preken-over-diverse-onderwerpen.docx",
        "Smijtegelt - 50 Keurstoffen"
    )
    return save_json(sermons, "smytegelt-50-keurstoffen.json")


def parse_smytegelt_deel(nr):
    """Parse 'Een Woord op zijn Tijd' delen 1-4."""
    sermons = split_sermons(
        f"smytegelt-deel-{nr}.docx",
        f"Smijtegelt - Een Woord op zijn Tijd deel {nr}",
        sermon_patterns=[
            # These use numbered ordinal + "PREEK/PREDIKATIE over"
            re.compile(r'^(EERSTE|TWEEDE|DERDE|VIERDE|VIJFDE|ZESDE|ZEVENDE|ACHTSTE|NEGENDE|TIENDE|'
                       r'ELFDE|TWAALFDE|DERTIENDE|VEERTIENDE|VIERTIENDE|VIJFTIENDE|ZESTIENDE|ZEVENTIENDE|'
                       r'ACHTTIENDE|NEGENTIENDE|TWINTIGSTE|EENENTWINTIGSTE|TWEE.NTWINTIGSTE|'
                       r'DRIE.NTWINTIGSTE|VIERENTWINTIGSTE|VIJFENTWINTIGSTE|ZESENTWINTIGSTE|'
                       r'ZEVENENTWINTIGSTE|ACHTENTWINTIGSTE|NEGENENTWINTIGSTE|DERTIGSTE|'
                       r'EENENDERTIGSTE|TWEE.NDERTIGSTE|DRIE.NDERTIGSTE|VIERENDERTIGSTE|'
                       r'VIJFENDERTIGSTE|ZESENDERTIGSTE|ZEVENENDERTIGSTE|ACHTENDERTIGSTE|'
                       r'NEGENENDERTIGSTE|VEERTIGSTE|EENENVEERTIGSTE|TWEE.NVEERTIGSTE|'
                       r'DRIE.NVEERTIGSTE|VIERENVEERTIGSTE|VIJFENVEERTIGSTE|ZESENVEERTIGSTE|'
                       r'ZEVENENVEERTIGSTE)\s+(preek|predikatie|PREEK|PREDIKATIE)',
                       re.IGNORECASE),
            # Heading style with "preek" in it
            re.compile(r'^\d+[eE]\s+(preek|predikatie|predicatie)', re.IGNORECASE),
        ]
    )
    return save_json(sermons, f"smytegelt-deel-{nr}.json")


def parse_smytegelt_zestal():
    sermons = split_sermons(
        "smytegelt-zestal-leerredenen.docx",
        "Smijtegelt - Zestal Leerredenen"
    )
    return save_json(sermons, "smytegelt-zestal-leerredenen.json")


def parse_smytegelt_52():
    """Parse 52 preken catechismus - splits on 'Zondag' pattern."""
    doc = Document(os.path.join(BASE, "smytegelt-52-preken-catechismus.docx"))
    paragraphs = extract_text(doc)

    source = "Smijtegelt - 52 Preken Catechismus"

    # These split on "Zondag X" or ordinal + predicatie/preek
    zondag_pat = re.compile(r'^(Eerste|Tweede|Derde|Vierde|Vijfde|Zesde|Zevende|Achtste|Negende|Tiende|'
                            r'Elfde|Twaalfde|Dertiende|Veertiende|Vijftiende|Zestiende|Zeventiende|'
                            r'Achttiende|Negentiende|Twintigste|Eenentwintigste|Twee.ntwintigste|'
                            r'Drie.ntwintigste|Vierentwintigste|Vijfentwintigste|Zesentwintigste|'
                            r'Zevenentwintigste|Achtentwintigste|Negenentwintigste|Dertigste|'
                            r'Eenendertigste|Twee.ndertigste|Drie.ndertigste|Vierendertigste|'
                            r'Vijfendertigste|Zesendertigste|Zevenendertigste|Achtendertigste|'
                            r'Negenendertigste|Veertigste|Eenenveertigste|Twee.nveertigste|'
                            r'Drie.nveertigste|Vierenveertigste|Vijfenveertigste|Zesenveertigste|'
                            r'Zevenenveertigste|Achtenveertigste|Negenenveertigste|Vijftigste|'
                            r'Eenenvijftigste|Twee.nvijftigste)\s+(predicatie|preek|predikatie)',
                            re.IGNORECASE)

    zondag_heading = re.compile(r'^Zondag\s+\d+', re.IGNORECASE)

    sermons = []
    current_title = None
    current_text = []
    in_preamble = True

    for text, style in paragraphs:
        is_sermon = zondag_pat.search(text) or zondag_heading.search(text)

        if is_sermon:
            in_preamble = False
            if current_title and current_text:
                sermons.append({
                    "title": current_title,
                    "text": "\n\n".join(current_text),
                    "source_collection": source
                })
            current_title = text
            current_text = []
        elif not in_preamble and current_title is not None:
            current_text.append(text)

    if current_title and current_text:
        sermons.append({
            "title": current_title,
            "text": "\n\n".join(current_text),
            "source_collection": source
        })

    return save_json(sermons, "smytegelt-52-preken-catechismus.json")


if __name__ == "__main__":
    total = 0

    print("=== Van der Groe: Toetssteen ===")
    for deel in [1, 2, 3]:
        total += parse_toetssteen(deel)

    print("\n=== Smijtegelt: 16 Predicaties ===")
    total += parse_smytegelt_16()

    print("\n=== Smijtegelt: 50 Keurstoffen ===")
    total += parse_smytegelt_50()

    print("\n=== Smijtegelt: Een Woord op zijn Tijd (4 delen) ===")
    for deel in [1, 2, 3, 4]:
        total += parse_smytegelt_deel(deel)

    print("\n=== Smijtegelt: Zestal Leerredenen ===")
    total += parse_smytegelt_zestal()

    print("\n=== Smijtegelt: 52 Preken Catechismus ===")
    total += parse_smytegelt_52()

    print(f"\n{'='*50}")
    print(f"TOTAAL: {total} items geparsed")
